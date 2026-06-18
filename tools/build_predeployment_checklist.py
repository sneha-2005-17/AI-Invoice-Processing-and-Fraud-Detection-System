from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\AI Invoice Processing and Fraud Detection System")
OUT = ROOT / "outputs" / "AI_Model_Pre_Deployment_Checklist_InvoiceAI.docx"
PORTFOLIO_IMAGE = Path(r"C:\Users\DELL\Desktop\portfolio and architecture.png")


ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(88, 88, 88)
RED = RGBColor(155, 28, 28)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, color=ACCENT)
        set_cell_text(cells[1], value)
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(4.7)


def add_status_line(doc: Document, status: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Status: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    s = p.add_run(status)
    s.bold = True
    if status in {"FAIL", "NOT APPROVED FOR DEPLOYMENT"}:
        s.font.color.rgb = RED
    elif status == "PASS":
        s.font.color.rgb = RGBColor(0, 110, 70)
    elif status == "PARTIAL":
        s.font.color.rgb = RGBColor(122, 90, 0)


def add_bullets(doc: Document, title: str, items: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{title}:")
    r.bold = True
    r.font.color.rgb = ACCENT
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        bp.add_run(item)


def add_section(
    doc: Document,
    number: int,
    title: str,
    evidence: list[str],
    assessment: list[str],
    status: str,
    recommendations: list[str],
    metrics: list[str] | None = None,
) -> None:
    doc.add_heading(f"{number}. {title}", level=1)
    add_bullets(doc, "Evidence Found", evidence or ["Evidence Not Found"])
    if metrics is not None:
        add_bullets(doc, "Metrics Found", metrics or ["Evidence Not Found"])
    add_bullets(doc, "Assessment", assessment)
    add_status_line(doc, status)
    add_bullets(doc, "Recommendations", recommendations)


def add_matrix(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    doc.add_heading("Deployment Readiness Matrix", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Area", "Status", "Risk Level", "Notes"]
    widths = [1.65, 1.0, 1.0, 2.85]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4D78")
        table.rows[0].cells[i].width = Inches(widths[i])
    for area, status, risk, notes in rows:
        cells = table.add_row().cells
        for i, text in enumerate([area, status, risk, notes]):
            set_cell_text(cells[i], text, bold=(i == 1))
            cells[i].width = Inches(widths[i])
        if status == "FAIL":
            set_cell_shading(cells[1], "F8D7DA")
        elif status == "PARTIAL":
            set_cell_shading(cells[1], "FFF3CD")
        elif status == "PASS":
            set_cell_shading(cells[1], "D1E7DD")


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11.5, ACCENT),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("AI Model Pre-Deployment Checklist")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("InvoiceAI Command - AI Invoice Processing and Fraud Detection System")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    add_kv_table(
        doc,
        [
            ("Project Name", "AI Invoice Processing and Fraud Detection System"),
            ("Repository", str(ROOT)),
            ("Version", "1.0.0 from FastAPI metadata and frontend package.json; Git metadata not present in workspace"),
            ("Analysis Date", date(2026, 6, 8).isoformat()),
            ("Architecture", "User -> Frontend (React/Next.js) -> FastAPI Backend -> Authentication Service -> Database -> LLM Service"),
        ],
    )

    doc.add_paragraph()
    if PORTFOLIO_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(PORTFOLIO_IMAGE), width=Inches(6.5))
        cap = doc.add_paragraph("Portfolio and conceptual system architecture supplied for the report.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(8.5)

    doc.add_page_break()


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    doc.add_heading("Executive Summary", level=1)
    add_kv_table(
        doc,
        [
            ("Overall Deployment Readiness Score", "43 / 100"),
            ("Deployment Recommendation", "NOT APPROVED FOR DEPLOYMENT"),
            ("Critical Risks", "Default admin credentials, placeholder secrets, unvalidated static evaluation metrics, no CI/CD test gate, no documented privacy/compliance controls, no drift/retraining plan."),
            ("Blocking Issues", "Remove demo credentials and secret defaults; add validated evaluation datasets/results; add backend test execution; implement production monitoring, migrations, backups, and privacy controls."),
        ],
    )
    doc.add_paragraph(
        "This review found a functional prototype with meaningful architecture components: FastAPI backend, Next.js/React frontend, JWT authentication, RBAC dependencies, OCR extraction, deterministic fraud scoring, ChromaDB-backed RAG, audit-log tables, Docker/Render deployment files, and a dashboard UI. The repository does not yet provide sufficient evidence for production deployment of an AI-assisted financial workflow. Several controls exist, but key operational, security, compliance, model governance, and validation artifacts are missing or only represented by static/demo values."
    )

    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph("Requested conceptual flow:")
    for step in [
        "User",
        "Frontend (React/Next.js)",
        "FastAPI Backend",
        "Authentication Service",
        "Database",
        "LLM Service",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "Repository evidence aligns with this flow through frontend/services/api.ts, backend/app/main.py router registration, backend/app/routers/auth.py authentication endpoints, backend/app/core/database.py SQLAlchemy configuration, and Gemini/Chroma integrations in invoice_service.py and policy_assistant.py."
    )

    sections = [
        (
            "Problem and Goal Validation",
            [
                "README.md states the project processes invoices with OCR extraction, fraud scoring, RAG finance policy assistance, LangGraph workflows, dashboards, RBAC, evaluation metrics, and deployment configuration.",
                "backend/API.md documents upload, fraud reanalysis, RAG chat, dashboard, evaluation, reports, and admin endpoints.",
                "frontend/app/page.tsx implements dashboard, invoice queue, policy assistant, upload, document chat, and demo-login workflows.",
            ],
            [
                "The intended business goal is clear: invoice ingestion, extraction, fraud review, and policy assistance.",
                "Success criteria are not formally defined beyond static metric examples; acceptance thresholds and stakeholder approval evidence are missing.",
            ],
            "PARTIAL",
            [
                "Add a product requirements document with measurable deployment objectives, target users, risk tolerance, and acceptance thresholds.",
                "Separate demo-mode goals from production goals, especially for financial decisions and human review.",
            ],
        ),
        (
            "Data Quality Check",
            [
                "OCRService extracts PDF text with pdfplumber and falls back to pytesseract over rendered pages; images use pytesseract then EasyOCR if available (backend/app/services/ocr_service.py).",
                "Invoice upload allows PDF, PNG, JPG/JPEG, TIFF, and BMP with a 15 MB limit (backend/app/services/invoice_service.py:11-26).",
                "Local sample uploads and ChromaDB files exist under backend/storage, but no labeled dataset, data dictionary, or data-quality report is present.",
            ],
            [
                "Basic file type and size controls exist, and OCR parsing is defensive.",
                "Evidence Not Found for training/validation dataset lineage, duplicate analysis, OCR quality sampling, ground-truth labels, data completeness checks, or PII classification.",
            ],
            "PARTIAL",
            [
                "Create a representative invoice validation set with ground truth for invoice number, vendor, GST, dates, totals, tax, line items, and fraud labels.",
                "Add automated data-quality reports for OCR confidence, parse failures, missing fields, duplicate invoices, and distribution drift.",
            ],
        ),
        (
            "Model Performance Evaluation",
            [
                "EvaluationService returns RAG, system, performance, and user metrics (backend/app/services/evaluation_service.py).",
                "frontend/app/page.tsx displays faithfulness, extraction F1, fraud recall, and processing latency values.",
                "backend/tests/test_ocr_service.py covers invoice-number parsing variants only.",
            ],
            [
                "RAG metrics: faithfulness 0.86, context precision 0.81, context recall 0.78, answer relevance 0.88.",
                "System metrics: fraud detection accuracy 0.84, precision 0.82, recall 0.79, F1 0.80.",
                "Performance metrics: latency 1450 ms, cost/request USD 0.012, API response time 220 ms.",
                "Verification result: frontend lint passed with five warnings; backend tests could not run because pytest is not installed.",
            ],
            [
                "Metrics exist in code, but no evaluation dataset, experiment output, timestamp, statistical method, or CI artifact validates them.",
                "Fraud scoring is rules-based, so performance should be measured against labeled historical or synthetic fraud cases.",
            ],
            "PARTIAL",
            [
                "Replace hard-coded evaluation values with reproducible benchmark outputs and store reports as artifacts.",
                "Add tests for OCR extraction, RAG citation accuracy, fraud rules, authorization, and upload validation.",
            ],
        ),
        (
            "Bias and Fairness Check",
            [
                "Evidence Not Found for fairness tests, protected-class analysis, vendor-region bias checks, threshold parity, or appeal workflows.",
                "Fraud rules include unknown/unapproved vendor and elevated vendor risk checks (backend/app/services/fraud_service.py).",
            ],
            [
                "Vendor approval and historical risk can create disparate impacts on new or smaller vendors if not governed.",
                "No fairness criteria, subgroup reports, or manual override governance are documented.",
            ],
            "FAIL",
            [
                "Define fairness risks for vendor approval, geography, business size, invoice amount, and supplier tenure.",
                "Add periodic false-positive/false-negative review by vendor segment and require human approval for high-risk actions.",
            ],
        ),
        (
            "Security Check",
            [
                "JWT tokens are created with subject, role, and expiry (backend/app/core/security.py:30-36); current users are decoded and loaded from DB (backend/app/dependencies.py).",
                "RBAC dependency require_roles protects admin, fraud reanalysis, dashboard, evaluation, and policy upload routes.",
                "SlowAPI rate limiting is configured globally at settings.rate_limit_per_minute (backend/app/main.py).",
                "Security headers set X-Content-Type-Options, X-Frame-Options, and X-XSS-Protection (backend/app/main.py).",
                "Prompt-injection marker blocking and bleach sanitization exist (backend/app/core/security.py:18-46).",
                "Critical risk: default secret_key fallback is change-me and .env.example uses change-me-in-production; startup seeds admin@example.com/password123 (backend/app/main.py).",
                "Frontend stores JWT in localStorage (frontend/app/page.tsx), increasing token exposure risk under XSS.",
            ],
            [
                "Authentication and authorization exist, but production security posture is not acceptable with default credentials and fallback secrets.",
                "Prompt injection protection is marker-based and narrow; RAG context is constrained but not fully defended against malicious uploaded documents.",
                "Upload validation checks extension and size, but MIME sniffing, antivirus/malware scanning, and storage isolation are not evidenced.",
            ],
            "PARTIAL",
            [
                "Remove seeded demo admin credentials from production startup and require first-admin provisioning.",
                "Fail startup when SECRET_KEY or DATABASE_URL are default placeholders in production.",
                "Move tokens to secure httpOnly cookies or add hardened token lifecycle controls.",
                "Add MIME validation, malware scanning, secure object storage, per-route rate limits, structured security logging, and redaction tests.",
            ],
        ),
        (
            "Privacy and Compliance Check",
            [
                "Database models store users, invoices, raw OCR text, uploaded files, audit logs, RAG chunks, feedback, and reports (backend/app/models/entities.py).",
                "mask_sensitive helper exists but is not visibly applied across logging/reporting paths (backend/app/core/security.py).",
                "Evidence Not Found for retention policies, consent, data subject rights, encryption at rest, PII inventory, DPIA, SOC2/ISO controls, or audit retention.",
            ],
            [
                "The system processes financial documents and potentially personal/vendor data but lacks documented privacy controls.",
                "Raw OCR text and uploaded files are persisted locally without documented retention, classification, or deletion policy beyond manual invoice deletion.",
            ],
            "FAIL",
            [
                "Create a privacy impact assessment and data retention schedule for uploads, OCR text, reports, embeddings, and logs.",
                "Add encryption-at-rest requirements, access logging, PII redaction, backup retention, and deletion workflows.",
            ],
        ),
        (
            "Scalability Check",
            [
                "Docker Compose defines PostgreSQL, ChromaDB, backend, and frontend services with volumes.",
                "backend/core/database.py uses SQLAlchemy pool_pre_ping, and deployment notes recommend managed PostgreSQL and persistent vector storage.",
                "DEPLOYMENT.md recommends object storage for uploads/reports and Alembic migrations before long-lived production schema evolution.",
                "Evidence Not Found for caching, background workers, queues, autoscaling policy, resource limits, or load-test results.",
            ],
            [
                "The architecture is service-separated and containerized, but long OCR/LLM work appears synchronous in request paths.",
                "Local filesystem persistence and synchronous processing are likely bottlenecks under concurrent load.",
            ],
            "PARTIAL",
            [
                "Introduce background jobs for OCR, LLM calls, report generation, and policy ingestion.",
                "Add Redis or equivalent caching where appropriate, resource limits, load testing, and object storage.",
            ],
        ),
        (
            "Latency and Response Time Check",
            [
                "EvaluationService includes static performance metrics: latency_ms 1450 and api_response_time_ms 220.",
                "DEPLOYMENT.md says to configure monitoring for API latency, OCR latency, Gemini spend, and fraud review outcomes.",
                "Evidence Not Found for measured benchmark logs, tracing, SLOs, or percentile latency dashboards.",
            ],
            [
                "Latency targets are not backed by reproducible test evidence.",
                "OCR, PDF conversion, EasyOCR/Tesseract, Chroma, and Gemini calls can exceed interactive response targets.",
            ],
            "PARTIAL",
            [
                "Define SLOs and measure p50/p95/p99 latencies for upload, OCR, fraud scoring, RAG chat, and dashboard APIs.",
                "Move long-running processing to asynchronous jobs with user-visible status.",
            ],
        ),
        (
            "Robustness and Edge Case Testing",
            [
                "OCR parsing has defensive try/except logic and tests invoice-number variants (backend/tests/test_ocr_service.py).",
                "Policy and invoice chat have fallback behavior when Gemini or Chroma is unavailable.",
                "Frontend lint passed with warnings only; backend tests could not run because pytest is absent.",
            ],
            [
                "There is limited automated robustness coverage for malformed PDFs, empty OCR, huge pages, wrong MIME, concurrent uploads, DB failures, and authorization bypass attempts.",
                "Several exception handlers swallow failures and return fallback content, which can obscure operational issues.",
            ],
            "PARTIAL",
            [
                "Add negative and edge-case API tests, file fuzzing cases, LLM failure tests, Chroma outage tests, and database transaction rollback tests.",
                "Log fallback activation with safe structured events and alert when fallbacks exceed thresholds.",
            ],
        ),
        (
            "Explainability and Output Validation",
            [
                "FraudDetectionService stores flags and explanation text per invoice (backend/app/services/fraud_service.py).",
                "RAG responses include citations, source document, section reference, and confidence score (backend/app/rag/policy_assistant.py).",
                "Invoice chat prompt instructs the model to answer strictly from document text and say when the answer is not found (backend/app/routers/invoices.py).",
            ],
            [
                "Fraud outputs are relatively explainable because they are rule-based and flag-driven.",
                "LLM output validation is basic: JSON parsing is attempted, but schema-level enforcement, hallucination tests, and citation faithfulness validation are not evidenced.",
            ],
            "PARTIAL",
            [
                "Add output schemas with strict validation and refusal/fallback paths for malformed LLM output.",
                "Add citation faithfulness tests and require high-risk fraud explanations to show rule evidence.",
            ],
        ),
        (
            "Integration Testing",
            [
                "frontend/services/api.ts centralizes API calls and bearer-token headers.",
                "Frontend calls /api/auth/login, /api/dashboard/summary, /api/invoices, /api/invoices/upload, /api/rag/chat, /api/rag/documents, and document chat routes.",
                "Docker Compose wires frontend, backend, PostgreSQL, and ChromaDB.",
                "Evidence Not Found for end-to-end tests, Playwright tests, API contract tests, or auth-flow integration tests.",
            ],
            [
                "Implementation paths align between frontend and backend for core flows.",
                "No automated integration evidence confirms frontend/backend/database/external API flows work under CI or production-like data.",
            ],
            "PARTIAL",
            [
                "Add Playwright or Cypress tests for login, upload, invoice list, fraud details, RAG chat, and admin policy upload.",
                "Add backend API tests using a test database and contract checks for frontend response types.",
            ],
        ),
        (
            "Monitoring and Logging Setup",
            [
                "AuditLog model and write_audit service exist; auth register/login and invoice upload write audit events.",
                "Admin audit-log route returns recent audit events for administrators.",
                "DEPLOYMENT.md recommends latency, OCR, Gemini spend, and fraud review outcome monitoring.",
                "Evidence Not Found for OpenTelemetry, metrics exporters, alert rules, centralized logs, dashboards, or log redaction tests.",
            ],
            [
                "Audit logging is a useful starting point, but production observability is not implemented.",
                "Operational metrics and security event monitoring are recommendations only, not repository artifacts.",
            ],
            "PARTIAL",
            [
                "Add structured logs, request IDs, error budgets, metrics exporters, alert rules, dashboard definitions, and secure log retention.",
                "Expand audit events to failed login, authorization denial, report download, policy upload, delete, and admin actions.",
            ],
        ),
        (
            "Model Drift and Retraining Plan",
            [
                "Evidence Not Found for drift detection, data distribution monitoring, retraining triggers, model cards, scheduled evaluation, or champion/challenger process.",
                "Fraud logic is deterministic rules; RAG behavior depends on uploaded policies, embeddings, Chroma retrieval, and Gemini model behavior.",
            ],
            [
                "No governance exists for changing invoice formats, vendor behavior, fraud patterns, policy updates, or external LLM version changes.",
                "Even rules-based systems need drift monitoring on false positives, false negatives, OCR parse rates, and RAG retrieval quality.",
            ],
            "FAIL",
            [
                "Define monitoring for OCR extraction drift, fraud flag rates, vendor-risk calibration, RAG retrieval metrics, and user feedback.",
                "Create a scheduled review and retraining/recalibration process with approval gates.",
            ],
        ),
        (
            "Version Control and Model Registry",
            [
                "README.md and package metadata identify version 1.0.0; backend FastAPI metadata also declares version 1.0.0.",
                "No .git directory is present in the workspace, and git CLI is unavailable in this environment.",
                "Evidence Not Found for release tags, changelog, model registry, prompt registry, artifact storage, experiment tracking, or deployment approvals.",
            ],
            [
                "The code has version labels but no verifiable repository history or release governance in the provided folder.",
                "Gemini model name and prompts are in code, but prompt/version lineage is not tracked as deployable artifacts.",
            ],
            "FAIL",
            [
                "Use Git with protected branches, release tags, changelog, and pull-request review requirements.",
                "Register prompts, rule versions, evaluation artifacts, RAG embeddings/index versions, and external model versions.",
            ],
        ),
        (
            "Deployment Environment Check",
            [
                "backend/Dockerfile and frontend/Dockerfile exist.",
                "docker-compose.yml defines PostgreSQL, ChromaDB, backend, and frontend with local volumes.",
                "backend/render.yaml defines a Render web service with DATABASE_URL, SECRET_KEY, GEMINI_API_KEY, and ALLOWED_ORIGINS.",
                "backend/.env.example and frontend/.env.example define required variables.",
                "Evidence Not Found for Kubernetes, SSL/TLS config, domain config, backup scripts, secret manager configuration, Alembic migrations, or production object storage config.",
            ],
            [
                "Container and Render deployment foundations exist.",
                "Production environment hardening is incomplete, especially migrations, backups, TLS/domain evidence, secret management, and persistent upload/report storage.",
            ],
            "PARTIAL",
            [
                "Add production deployment runbook, migration tooling, backup/restore verification, domain/TLS configuration, and secret-manager integration.",
                "Avoid using .env.example as a Docker Compose env_file with placeholder secrets for anything beyond local demo.",
            ],
        ),
        (
            "Fallback and Error Handling",
            [
                "OCR returns empty text on extraction failure rather than mocked invoice content (backend/app/services/ocr_service.py).",
                "PolicyAssistant falls back to no-context or context excerpt when Chroma/Gemini fails.",
                "InvoiceService has fallback document analysis when Gemini is unavailable.",
                "Report routes convert service errors to HTTP 404/503.",
            ],
            [
                "Fallbacks prevent crashes, but they may silently degrade AI quality and confidence.",
                "No global error schema, retry strategy, circuit breaker, dead-letter handling, or user-facing processing status model is evidenced.",
            ],
            "PARTIAL",
            [
                "Add explicit degraded-mode status, error codes, retry/backoff policies, and alerting for fallback activation.",
                "Use asynchronous processing states for upload/OCR/LLM operations and expose failure reasons safely.",
            ],
        ),
        (
            "Human Review and Approval",
            [
                "Frontend has invoice queue/review UI concepts and fraud result display.",
                "Fraud reanalysis is restricted to admin/analyst roles.",
                "Evidence Not Found for formal approval workflow, reviewer assignment, maker-checker controls, sign-off logs, escalation policy, or override reason capture.",
            ],
            [
                "High-risk financial decisions require explicit human review controls; the repository currently presents review information but does not enforce approval gates.",
                "No evidence confirms management, compliance, or security approval before deployment.",
            ],
            "FAIL",
            [
                "Implement review states, assigned reviewer, approval/deny actions, override reason, escalation thresholds, and immutable approval audit logs.",
                "Require sign-off from product owner, security, compliance/privacy, and operations before production release.",
            ],
        ),
    ]

    for idx, item in enumerate(sections, start=1):
        if idx == 3:
            add_section(doc, idx, item[0], item[1], item[3], item[4], item[5], metrics=item[2])
        else:
            add_section(doc, idx, item[0], item[1], item[2], item[3], item[4])

    matrix_rows = [
        ("Problem and Goal Validation", "PARTIAL", "Medium", "Goal is clear; success criteria and approvals missing."),
        ("Data Quality", "PARTIAL", "High", "OCR controls exist; labeled datasets and quality reports missing."),
        ("Model Performance", "PARTIAL", "High", "Static metrics found; reproducible evaluation missing."),
        ("Bias and Fairness", "FAIL", "High", "No fairness tests or subgroup analysis."),
        ("Security", "PARTIAL", "Critical", "JWT/RBAC exists; default credentials and placeholder secrets block deployment."),
        ("Privacy and Compliance", "FAIL", "Critical", "Financial data persisted without documented privacy controls."),
        ("Scalability", "PARTIAL", "Medium", "Containerized services; synchronous OCR/LLM and filesystem storage risk."),
        ("Latency", "PARTIAL", "Medium", "Static latency values; no benchmark evidence."),
        ("Robustness", "PARTIAL", "High", "Limited tests; backend test runner missing."),
        ("Explainability", "PARTIAL", "Medium", "Rule flags and citations exist; LLM validation incomplete."),
        ("Integration Testing", "PARTIAL", "High", "API paths align; no E2E or contract tests."),
        ("Monitoring and Logging", "PARTIAL", "High", "Audit logs exist; observability stack missing."),
        ("Drift and Retraining", "FAIL", "High", "No drift/retraining governance."),
        ("Version Control and Registry", "FAIL", "High", "No .git or registry evidence in provided folder."),
        ("Deployment Environment", "PARTIAL", "High", "Docker/Render present; backups, TLS, migrations missing."),
        ("Fallback/Error Handling", "PARTIAL", "Medium", "Fallbacks exist; degraded-mode governance missing."),
        ("Human Review", "FAIL", "Critical", "No enforced approval workflow or sign-off evidence."),
    ]
    add_matrix(doc, matrix_rows)

    doc.add_heading("Missing Evidence", level=1)
    for item in [
        "Validated production requirements, success thresholds, and stakeholder sign-off.",
        "Representative labeled evaluation datasets for OCR, extraction, fraud, RAG, and document chat.",
        "Bias/fairness test reports and subgroup false-positive/false-negative analysis.",
        "CI/CD pipelines, protected branches, release tags, changelog, and Git history in the provided folder.",
        "Backend test execution environment with pytest installed; end-to-end and API contract tests.",
        "Model registry, prompt registry, artifact storage, RAG index versioning, and experiment tracking.",
        "Privacy impact assessment, retention schedule, PII inventory, encryption-at-rest evidence, and compliance mapping.",
        "Production monitoring dashboards, alert rules, tracing, structured log retention, and log redaction verification.",
        "Load tests, latency benchmarks, autoscaling/resource policies, caching strategy, and background job architecture.",
        "Database migration tooling, backup/restore procedures, TLS/domain configuration, and object storage setup.",
        "Formal human review workflow, approval audit trail, override process, and deployment approval record.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Critical Findings", level=1)
    for item in [
        "Default admin@example.com/password123 account is seeded during startup and demo credentials are prefilled in the frontend; this must not ship to production.",
        "Secret defaults and examples contain placeholder values such as change-me and change-me-in-production; production startup should fail closed if placeholders are used.",
        "Evaluation metrics are hard-coded/static and not tied to a reproducible benchmark, dataset, experiment, or CI artifact.",
        "Privacy/compliance evidence is insufficient for a system storing financial documents, raw OCR text, reports, embeddings, and audit logs.",
        "No CI/CD pipeline, backend test execution, model/prompt registry, drift monitoring, or formal approval workflow was found.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Final Recommendation", level=1)
    add_status_line(doc, "NOT APPROVED FOR DEPLOYMENT")
    doc.add_paragraph(
        "Based only on artifacts present in the project folder, the system should not be deployed to production yet. The project is a strong prototype with visible end-to-end functionality, but production deployment of an AI-enabled financial workflow requires validated metrics, hardened secrets and authentication, privacy controls, monitoring, reproducible testing, release governance, and human approval controls. The recommended path is to approve continued development and controlled demonstration only, then reassess after the critical findings are remediated and independently verified."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "AI Model Pre-Deployment Checklist - InvoiceAI Command - Evidence-based repository review"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = MUTED

    doc.save(OUT)


if __name__ == "__main__":
    build()
